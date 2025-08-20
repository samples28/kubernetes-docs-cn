#!/usr/bin/env python3
"""
将Markdown文件转换为Word文档的脚本
需要安装: pip install python-docx markdown beautifulsoup4
"""

import os
import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import re

def markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换Markdown为HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Inches(0.12)
    
    # 处理HTML元素
    for element in soup.find_all():
        if element.name == 'h1':
            heading = doc.add_heading(element.get_text(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            if element.get_text().strip():
                doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'table':
            # 处理表格
            rows = element.find_all('tr')
            if rows:
                # 计算列数
                max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        if j < max_cols:
                            table.cell(i, j).text = cell.get_text().strip()
        elif element.name == 'blockquote':
            doc.add_paragraph(element.get_text(), style='Quote')
        elif element.name == 'code':
            doc.add_paragraph(element.get_text(), style='Intense Quote')
    
    # 保存文档
    doc.save(docx_file)
    print(f"已转换: {md_file} -> {docx_file}")

def main():
    """主函数"""
    # 要转换的文件列表
    files_to_convert = [
        "技术更新管理政策.md",
        "系统清单和版本管理记录.md", 
        "CTO审批记录模板.md",
        "技术更新执行报告.md",
        "安全补丁管理流程.md"
    ]
    
    # 创建输出目录
    output_dir = "word_documents"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 转换每个文件
    for md_file in files_to_convert:
        if os.path.exists(md_file):
            docx_file = os.path.join(output_dir, md_file.replace('.md', '.docx'))
            try:
                markdown_to_docx(md_file, docx_file)
            except Exception as e:
                print(f"转换失败 {md_file}: {e}")
        else:
            print(f"文件不存在: {md_file}")
    
    print(f"\n所有文档已转换完成，保存在 {output_dir} 目录中")

if __name__ == "__main__":
    main()
